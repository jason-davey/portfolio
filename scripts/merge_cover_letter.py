#!/usr/bin/env python3
"""
Cover Letter Template Merger
Merges markdown cover letter content into a Word template while preserving template structure
"""

import sys
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

def parse_markdown_cover_letter(md_file):
    """Parse markdown cover letter and extract key components"""
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Extract sections
    data = {
        'name': '',
        'title': '',
        'email': '',
        'phone': '',
        'linkedin': '',
        'portfolio': '',
        'location': '',
        'job_role': '',
        'company': '',
        'greeting': 'Dear Hiring Manager,',
        'body': '',
        'closing': 'Sincerely,',
    }

    # Parse header section
    lines = content.split('\n')
    in_header = False
    body_started = False
    body_lines = []

    for i, line in enumerate(lines):
        line = line.strip()

        # Extract name (first # heading)
        if line.startswith('# ') and not data['name']:
            data['name'] = line.replace('# ', '').strip()
            continue

        # Extract title (bold text after name)
        if line.startswith('**') and '|' in line and not data['title']:
            data['title'] = line.replace('**', '').strip()
            continue

        # Extract contact info
        if '@' in line and not data['email']:
            # Parse contact line: email | phone
            parts = [p.strip() for p in line.split('|')]
            for part in parts:
                if '@' in part:
                    data['email'] = part
                elif '+' in part or part.replace(' ', '').replace('-', '').isdigit():
                    data['phone'] = part

        # Extract linkedin and portfolio
        if 'linkedin.com' in line.lower():
            match = re.search(r'linkedin\.com/[^\s|]+', line)
            if match:
                data['linkedin'] = match.group(0)

        if '.myportfolio.com' in line.lower() or 'portfolio' in line.lower():
            parts = line.split('|')
            for part in parts:
                if 'portfolio' in part.lower() and 'linkedin' not in part.lower():
                    data['portfolio'] = part.strip()

        # Extract location
        if any(state in line for state in ['NSW', 'VIC', 'QLD', 'SA', 'WA', 'TAS', 'NT', 'ACT']):
            data['location'] = line
            continue

        # Extract job role from ## heading
        if line.startswith('## '):
            role = line.replace('## ', '').strip()
            if 'Dear' not in role and not body_started:
                data['job_role'] = role
            continue

        # Extract greeting
        if line.startswith('Dear '):
            data['greeting'] = line
            body_started = True
            continue

        # Collect body paragraphs
        if body_started and line and not line.startswith('#'):
            # Skip the "Sincerely" and signature parts
            if line.startswith('Sincerely') or line.startswith('**') and line.endswith('**') and len(line) < 50:
                data['closing'] = line.replace('**', '').strip()
                break

            # Skip metadata lines
            if line.startswith('*Enclosure:'):
                continue

            if line.startswith('---'):
                continue

            body_lines.append(line)

    # Join body paragraphs
    data['body'] = '\n\n'.join([line for line in body_lines if line])

    return data

def replace_placeholder(doc, placeholder, replacement):
    """Replace placeholder text in document"""
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            # Replace inline
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, replacement)

def insert_body_content(doc, body_text, marker='[Cover letter details'):
    """Insert body content at marker position"""
    for i, paragraph in enumerate(doc.paragraphs):
        if marker in paragraph.text:
            # Clear the marker paragraph
            paragraph.clear()

            # Split body into paragraphs and insert
            paragraphs = body_text.split('\n\n')

            # Insert first paragraph in place of marker
            if paragraphs:
                paragraph.text = paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                # Insert remaining paragraphs after
                for para_text in paragraphs[1:]:
                    if para_text.strip():
                        new_para = paragraph.insert_paragraph_before(para_text)
                        new_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            return True
    return False

def merge_cover_letter(template_path, md_file, output_path):
    """Main merge function"""

    print(f"📄 Reading template: {template_path}")

    # Handle .dotx template files by converting to .docx first
    if template_path.endswith('.dotx') or template_path.endswith('.dot'):
        import subprocess
        import tempfile

        print("🔄 Converting template to .docx format...")
        temp_docx = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        temp_docx.close()

        # Use LibreOffice to convert template to docx
        result = subprocess.run([
            'soffice', '--headless', '--convert-to', 'docx',
            '--outdir', os.path.dirname(temp_docx.name),
            template_path
        ], capture_output=True, text=True)

        # Find the converted file
        template_base = os.path.splitext(os.path.basename(template_path))[0]
        converted_file = os.path.join(os.path.dirname(temp_docx.name), f"{template_base}.docx")

        if os.path.exists(converted_file):
            os.rename(converted_file, temp_docx.name)
            template_path = temp_docx.name
        else:
            print(f"⚠️  Warning: Could not convert template, using as-is")

    try:
        doc = Document(template_path)
    except Exception as e:
        print(f"❌ Error opening template: {e}")
        print("Trying with default document...")
        doc = Document()

    print(f"📝 Parsing markdown: {md_file}")
    data = parse_markdown_cover_letter(md_file)

    print("🔄 Merging content into template...")

    # Replace placeholders
    replace_placeholder(doc, '[HIRING COMPANY]', data.get('company', '[HIRING COMPANY]'))
    replace_placeholder(doc, '[HIRING COMPANY MANAGER]', data.get('greeting', 'Dear Hiring Manager'))
    replace_placeholder(doc, 'OPPORTUNITY / JOB ROLE', data.get('job_role', 'OPPORTUNITY / JOB ROLE'))
    replace_placeholder(doc, '[HIRING COMPANY ADDRESS]', '')

    # Update greeting
    if data['greeting']:
        replace_placeholder(doc, 'Dear [HIRING COMPANY MANAGER],', data['greeting'])

    # Insert body content
    if not insert_body_content(doc, data['body']):
        print("⚠️  Warning: Could not find body content marker")

    # Save
    print(f"💾 Saving to: {output_path}")
    doc.save(output_path)
    print("✅ Merge completed successfully!")

if __name__ == '__main__':
    if len(sys.argv) < 4:
        print("Usage: python3 merge_cover_letter.py <template.dotx> <input.md> <output.docx>")
        sys.exit(1)

    template_path = sys.argv[1]
    md_file = sys.argv[2]
    output_path = sys.argv[3]

    if not os.path.exists(template_path):
        print(f"❌ Error: Template not found: {template_path}")
        sys.exit(1)

    if not os.path.exists(md_file):
        print(f"❌ Error: Markdown file not found: {md_file}")
        sys.exit(1)

    merge_cover_letter(template_path, md_file, output_path)
